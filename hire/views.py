from django.shortcuts import render
import requests
from bs4 import BeautifulSoup
import openai
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from docx import Document
import re
import os
from docx import Document
import re
from nltk.stem import WordNetLemmatizer
from django.http import HttpResponse
from django.http import Http404
from django.http import HttpResponse

# Create your views here.
def home(request):
    return render(request,'index.html')
def jobseeker(request):
    if request.method == 'POST':
        search_query = request.GET.get('q', '') 
    
        job_tech =[]
        job1=[]
        job2=[]
        type_value = " "
    

        # openai.api_key = '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'
        openai.api_key = os.environ.get('OPENAI_API_KEY')

        response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=f"Get the Job Role , Education , Place , Technologies & Skills , Experience , Job Type from Description:\n{search_query}\n",
        
        temperature=0.5,
        max_tokens=100,
        n=1,
        stop=None,
        timeout=10,
        
        )

        main_summary = response.choices[0].text.strip()

        print(main_summary)

        if main_summary != " ":       
            # Extract specific information using regular expressions
            job_role = re.search(r'Job Role:(.*)', main_summary).group(1).strip()
            education = re.search(r'Education:(.*)', main_summary).group(1).strip()
            place = re.search(r'Place:(.*)', main_summary).group(1).strip()
            technologies = re.search(r'Technologies & Skills:(.*)', main_summary).group(1).strip()
            experience = re.search(r'Experience:(.*)', main_summary).group(1).strip()
            # job_type = re.search(r'Job Type:(.*)', main_summary).group(1).strip()

            match = re.search(r'Job Type:(.*)', main_summary)
            if match:
                job_type = match.group(1).strip()
            else:
                job_type = "N/A"

            # Create a dictionary for the summary
            summary_dict = {
                'Job_Role': job_role,
                'Education': education,
                'Place': place,
                'Technologies': technologies,
                'Experience': experience,
                'Job_Type': job_type,
            }

            # Print the summary dictionary
            print(summary_dict)

            value1 = summary_dict.get("Job_Role")
            value2 = summary_dict.get("Education")
            value3 = summary_dict.get("Place")
            value4 = summary_dict.get("Technologies")
            value5 = summary_dict.get("Experience")
            value6 = summary_dict.get("Job_Type")



            if value1 != "N/A" and value1 !="Not Specified" and value1 != "None" and value1 != " ":
                role_values = word_tokenize(value1)
                stop_words =stopwords.words("english")
                stop_words.extend(['.',",",":",")",","])

                filtered_words = []
                for token in role_values:
                    if token not in stop_words:
                        filtered_words.append(token)
                
                old_char = '/'
                new_char = '-'

                updated_list = []

                for string in filtered_words:
                    updated_string = string.replace(old_char, new_char)
                    updated_list.append(updated_string)
                job_tech.append(updated_list)
            else:
                pass


            # if value2 != "N/A" and value2 !="Not Specified" and value2 != "None" and value2 != " ":
            #     job.append(value2)
            # else:
            #     pass


            if value3 != "N/A" and value3 !="Not Specified" and value3 != "None" and value3 != " " and value3 != "undefined" and value3 != "Unknown ":
                place_values = word_tokenize(value3.lower())
                stop_words =stopwords.words("english")
                stop_words.extend(['.',",",":",")",",","any"])

                filtered_words1 = []
                for token in place_values:
                    if token not in stop_words:
                        filtered_words1.append(token)
                
                old_char = '/'
                new_char = '-'
                updated_list1 = []

                for string in filtered_words1:
                    updated_string = string.replace(old_char, new_char)
                    updated_list1.append(updated_string)
                job1.append(updated_list1)
            else:
                pass


            if value4 != "N/A" and value4 !="Not Specified" and value4 != "None" and value4 != " ":
                tech_values = word_tokenize(value4)
                stop_words =stopwords.words("english")
                stop_words.extend(['.',",",":",")",",","("])

                filtered_words2 = []
                for token in tech_values:
                    if token not in stop_words:
                        filtered_words2.append(token)

                old_char = '/'
                new_char = '-'

                updated_list2 =[]
                for string in filtered_words2:
                    updated_string1 = string.replace(old_char, new_char)
                    updated_list2.append(updated_string1)
                job_tech.append(updated_list2)
            else:
                pass

            if value5 != "N/A" and value5 !="Not Specified" and value5 != "None" and value5 != " " and value5 !="0":

                pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"
            
                matches = re.findall(pattern, value5)
                
                numeric_matches = [match[0] for match in matches if match[0]]
                word_matches = [match[1] for match in matches if match[1]]
                
                totals = numeric_matches + word_matches

                print("lalitha")
                print(totals)
                print("lalitha")

                flattened3 = []
                for item in totals:
                    if isinstance(item, list):
                        flattened3.extend(item)
                    else:
                        flattened3.append(item)
                print(flattened3)
                string_result = str(flattened3)
                print(string_result)
                if string_result != " ":
                    if string_result != "0" and string_result < "3.1":
                        string_result = float(string_result)
                        string_result = string_result*12
                        exp=int(string_result)
                        print(exp)
                    else:
                        exp=500
            else:
                pass


            if value6 != "N/A" and value6 != "Not Specified" and value6 != "None" and value6 != " "and value6 != "None Specified":
                value6 = value6.lower()
                if value6 == "full time" or value6 == "Full-time":
                    type_value = 1
                elif value6 == "part time":
                    type_value = 2
                elif value6 == "internship":
                    type_value = 3
                elif value6 == "apprenticeship":
                    type_value = 6
                else:
                    type_value = 0
            else:
                pass
            # print(job)

            flattened = []
            for item in job_tech:
                if isinstance(item, list):
                    flattened.extend(item)
                else:
                    flattened.append(item)
            print(flattened)
            joined_string = '-'.join(str(element) for element in flattened)
            joined_string = joined_string.lower()
            print(joined_string)

            flattened1 = []
            for item in job1:
                if isinstance(item, list):
                    flattened1.extend(item)
                else:
                    flattened1.append(item)


            print(flattened1)
            joined_string1 = '-'.join(str(element) for element in flattened1)
            joined_string1 = joined_string1.lower()
            print(joined_string1)

            # if joined_string and joined_string1 and exp and type_value :
            #     url = "\n".join([f"https://www.freshersworld.com/jobs/jobsearch/{joined_string}-jobs-in-{joined_string1}?experience={exp}&job_type={type_value}".format(joined_string,joined_string1,string_result,type_value)])
            # if joined_string and joined_string1 and exp :
            #     url = "\n".join([f"https://www.freshersworld.com/jobs/jobsearch/{joined_string}-jobs-in-{joined_string1}?experience={exp}".format(joined_string,joined_string1,string_result)])
            if joined_string and joined_string1 :
                url = "\n".join([f"https://www.freshersworld.com/jobs/jobsearch/{joined_string}-jobs-in-{joined_string1}".format(joined_string,joined_string1)])
            elif joined_string :
                url = "\n".join([f"https://www.freshersworld.com/jobs/jobsearch/{joined_string}-jobs".format(joined_string)])
            else:
                print("Give some detailed information")
            print(url)
        else:
            print("Give some detailed information")  
            


        # url = "https://www.freshersworld.com/jobs/jobsearch/"+a[0]+"-"+a[1]+"-jobs-for-"+a[2]+"-in-"+a[3]+"?course=16"

        print(url)
        response = requests.get(url, allow_redirects=True)
        soup = BeautifulSoup(response.content, 'html.parser')
        job_listings = soup.find_all('div', class_='job-container')

        all_results = [] 

        for job in job_listings:
            job_title = job.find('span', class_='wrap-title').text
            company_name = job.find('h3', class_='latest-jobs-title font-16 margin-none inline-block company-name').text
            job_location = job.find('span', class_='job-location display-block modal-open job-details-span').text
            qualifications = job.find('span', class_='qualifications display-block modal-open pull-left job-details-span').text

            view_apply_button = job.find('span', class_='view-apply-button')
            view_apply_link = ''
            if view_apply_button:
                parent_div = view_apply_button.find_parent('div', class_='job-container')
                if parent_div:
                    view_apply_link = parent_div.get('job_display_url', '')

                job_title_formatted = job_title[:-4].rstrip()
                results = [job_title_formatted, company_name, job_location, qualifications, view_apply_link]
                all_results.append(results)  # Add the results to the list
        context = {
            'results': all_results,
            'search_query': search_query
        }
        return render(request, 'jobseeker.html', context)
    else:
        return render(request,'jobseeker.html')
def resume_finder(request):
    if request.method == 'POST':
        description = request.POST.get('description', '')
        

        
        matching_files = []
        matching_files1 = []
        matching_files2 = []

        # openai.api_key = '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'
        openai.api_key = os.environ.get('OPENAI_API_KEY')
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=f"Summarize in a short form and get the technologies:\n{description}\n",
            temperature=0.5,
            max_tokens=100,
            n=1,
            stop=None,
            timeout=10,
        )

        summary = response.choices[0].text.strip()
        # lemmatizer = WordNetLemmatizer()

        # lemmatized_words = [lemmatizer.lemmatize(word) for word in summary]
        # print(lemmatized_words)
        print(summary)
        stop_words =stopwords.words("english")
        stop_words.extend(['.',",",":",")",",","software","(","want","person","knowledge","skills","percent","years","year","different","some", "experience","expertise","technology","technologies", "get", "must","resume","resumes","developer", "latest","engineer","exp","yrs",";","tech","experience"])
        folder_path = "C:/Users/sp13/OneDrive/Desktop/docs"
        
        files = os.listdir(folder_path)
        for file in files:
            if file.endswith('.docx'):
                document = Document(os.path.join(folder_path, file))

                text = "\n".join([paragraph.text for paragraph in document.paragraphs])
                text = text.lower()

            
                pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"
                
                matches = re.findall(pattern, text)
                
                numeric_matches = [match[0] for match in matches if match[0]]
                word_matches = [match[1] for match in matches if match[1]]
                
                totals = numeric_matches + word_matches

                print(totals)


                summary_pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"

                matches = re.findall(summary_pattern, description)
                
                
                numeric_matches1 = [match[0] for match in matches if match[0]]
                word_matches1 = [match[1] for match in matches if match[1]]
                
                sum_totals = numeric_matches1 + word_matches1

                print(sum_totals)


                state = ""
                for sum_total in sum_totals:
                    for total in totals:
                        if sum_total in total:
                            state = str("Experience {} years match " .format(total))
                            print(state)
                    

                tokens = word_tokenize(summary.lower())
                filtered_words = []
                for token in tokens:
                    if token not in stop_words:
                        filtered_words.append(token)
                print(filtered_words)
                num_keywords = len(filtered_words)
                matching_keywords = []
                for word in filtered_words:
                    if word in text:
                        matching_keywords.append(word)



                matching_keywords1 = len(matching_keywords)
                print(matching_keywords)
                matching_percentage = matching_keywords1 / num_keywords * 100
                print(matching_percentage)
                matching = []

                if matching_percentage>=1:

                    if matching_percentage >= 75:
                        matching_files.append((file, matching_percentage,matching_keywords,state))
                    elif matching_percentage>=50 and matching_percentage<75:
                        matching_files1.append((file, matching_percentage,matching_keywords,state))
                    elif matching_percentage>=1 and matching_percentage<50:
                    # else:
                        matching_files2.append((file, matching_percentage,matching_keywords,state)) 
                    else:
                        print("No Files Found")

                else:
                    print("No Files Found")
        
        sorted_files1 = sorted(matching_files, key=lambda x: x[1], reverse=True)

        sorted_files2 = sorted(matching_files1, key=lambda x: x[1], reverse=True)

        sorted_files3 = sorted(matching_files2, key=lambda x: x[1], reverse=True)
        
        filtered_files=[]
        folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

        for file_info in sorted_files1:
            file_name = file_info[0]
            # file_path = folder_path + file_name
            filtered_files.append(file_name)

        print("filtered_files with sorted_files1 start")
        print(filtered_files)
        print("filtered_files with sorted_files1 end")




        filtered_files1=[]
        folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

        for file_info in sorted_files2:
            file_name = file_info[0]
            # file_path = folder_path + file_name
            filtered_files1.append(file_name)

        print("filtered_files1 with sorted_files2 start")
        print(filtered_files1)
        print("filtered_files1 with sorted_files2 end")

        filtered_files2=[]
        folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"


        resumes = []
        for file_info in sorted_files3:
            file_name = file_info[0]
            # file_path = folder_path + file_name

            filtered_files2.append(file_name)

        print("filtered_files2 with sorted_files3 start")
        print(filtered_files2)
        print("filtered_files2 with sorted_files3 start")
        

        matching_files_str = "<br><br>".join([f"{file}         {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state} " for file, percentage,matching_keywords,state in sorted_files1])
        matching_files_str1 = "<br><br>".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files2])
        matching_files_str2 = "<br><br>".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files3])
        

        count = len(matching_files + matching_files1 + matching_files2)

        print(f" {count} files were filtered ".format(count))

        context = {
            # 'results': resumes,
            'count':count,
            'files1':filtered_files,
            'matching_files_str':matching_files_str,
            'files2':filtered_files1,
            'matching_files_str1':matching_files_str1,
            'files3':filtered_files2,
            'matching_files_str2':matching_files_str2,
            'description': description
        }

        # return count,filtered_files,matching_files_str,filtered_files1,matching_files_str1,filtered_files2,matching_files_str2

        return render(request,'resume.html',context)
    else:
        return render(request,'resume.html')
    
def download_filtered_file(request, filename):
    file_path = os.path.join("C:/Users/sp13/OneDrive/Desktop/docs/", filename)
    with open(file_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type="application/octet-stream")
        response['Content-Disposition'] = f"attachment; filename={filename}"
        return response
    


def jobseeker_classic(request):
    if request.method == 'POST':
        skills = request.POST.get('skills', '')
        location = request.POST.get('location', '')
        experience = request.POST.get('experience', '')

        #skills condition 

        if experience != " ":

            pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"
            
            matches = re.findall(pattern, experience)
            
            numeric_matches = [match[0] for match in matches if match[0]]
            word_matches = [match[1] for match in matches if match[1]]
            
            totals = numeric_matches + word_matches

            print(totals)
            exp=1
            for i in totals:
                exp = int(i)
                if exp != 0 and exp < 3.1 :
                    exp = exp*12
                    exp=int(exp)
                    print(exp)
                else:
                    exp=500
        else:
            pass
        job_tech = []


        #skills condition 

        if skills != " ":
            skills = skills.lower()
            role_values = word_tokenize(skills)
            stop_words =stopwords.words("english")
            stop_words.extend(['.',",",":",")",","])

            filtered_words = []
            for token in role_values:
                if token not in stop_words:
                    filtered_words.append(token)
            
            old_char = '/'
            new_char = '-'

            updated_list = []

            for string in filtered_words:
                updated_string = string.replace(old_char, new_char)
                updated_list.append(updated_string)
            job_tech.append(updated_list)

            flattened1 = []
            for item in job_tech:
                if isinstance(item, list):
                    flattened1.extend(item)
                else:
                    flattened1.append(item)


            print(flattened1)
            joined_string1 = '-'.join(str(element) for element in flattened1)
            joined_string1 = joined_string1.lower()
        else:
            pass


        #location condition
        job_tech1 =[]
        if location != " ":
            location = location.lower()
            role_values = word_tokenize(location)
            stop_words =stopwords.words("english")
            stop_words.extend(['.',",",":",")",","])

            filtered_words1 = []
            for token in role_values:
                if token not in stop_words:
                    filtered_words1.append(token)
            
            old_char = '/'
            new_char = '-'

            updated_list1 = []

            for string in filtered_words1:
                updated_string1 = string.replace(old_char, new_char)
                updated_list1.append(updated_string1)
            job_tech1.append(updated_list1)

            flattened2 = []
            for item in job_tech1:
                if isinstance(item, list):
                    flattened2.extend(item)
                else:
                    flattened2.append(item)


            print(flattened2)
            joined_string2 = '-'.join(str(element) for element in flattened2)
            joined_string2 = joined_string2.lower()
        else:
            pass


        if skills and location and experience :
            url = "\n".join([f"https://www.freshersworld.com/jobs/jobsearch/{joined_string1}-jobs-in-{joined_string2}?experience={exp}".format(joined_string1,joined_string2,exp)])
        elif skills and location :
            url = "\n".join([f"https://www.freshersworld.com/jobs/jobsearch/{joined_string1}-jobs-in-{joined_string2}".format(joined_string1,joined_string2)])
        elif skills :
            url = "\n".join([f"https://www.freshersworld.com/jobs/jobsearch/{joined_string1}-jobs".format(joined_string1)])

        else:
            print("Give some detailed information")
        print(url)
        response = requests.get(url, allow_redirects=True)
        soup = BeautifulSoup(response.content, 'html.parser')
        job_listings = soup.find_all('div', class_='job-container')

        all_results = [] 

        for job in job_listings:
            job_title = job.find('span', class_='wrap-title').text
            company_name = job.find('h3', class_='latest-jobs-title font-16 margin-none inline-block company-name').text
            job_location = job.find('span', class_='job-location display-block modal-open job-details-span').text
            qualifications = job.find('span', class_='qualifications display-block modal-open pull-left job-details-span').text

            view_apply_button = job.find('span', class_='view-apply-button')
            view_apply_link = ''
            if view_apply_button:
                parent_div = view_apply_button.find_parent('div', class_='job-container')
                if parent_div:
                    view_apply_link = parent_div.get('job_display_url', '')
                job_title_formatted = job_title[:-4].rstrip()
                results = [job_title_formatted, company_name, job_location, qualifications, view_apply_link]
                all_results.append(results)  # Add the results to the list

        for result in all_results:
            print('Job Title:', result[0])
            print('Company Name:', result[1])
            print('Job Location:', result[2])
            print('Qualifications:', result[3])
            print('View & Apply Link:', result[4])
            print('---')   
        context = {
            'results': all_results,
            'skills': skills,
            'experience':experience,
            'location':location
        }
        return render(request, 'jobseeker-classic.html', context) 
    else:
            return render(request,'jobseeker-classic.html')
    
def resume_finder_classic(request):
    if request.method == 'POST':
        skills = request.POST.get('skills', '')
        experience = request.POST.get('experience', '')

        if skills and experience:

            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs"
            matching_files = []
            matching_files1 = []
            matching_files2 = []
            stop_words =stopwords.words("english")
            stop_words.extend(['.',",",":",")",",","software","(","want","person","knowledge","skills","percent","years","year","different","some", "experience","expertise","technology","technologies", "get", "must","resume","resumes","developer", "latest","engineer","exp","yrs",";","tech","experience"])
            files = os.listdir(folder_path)
            for file in files:
                if file.endswith('.docx'):
                    document = Document(os.path.join(folder_path, file))

                    text = "\n".join([paragraph.text for paragraph in document.paragraphs])
                    text = text.lower()


                    pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"
                    
                    matches = re.findall(pattern, text)
                    
                    numeric_matches = [match[0] for match in matches if match[0]]
                    word_matches = [match[1] for match in matches if match[1]]
                    
                    totals = numeric_matches + word_matches

                    print(totals)


                    summary_pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"

                    matches = re.findall(summary_pattern, experience)
                    
                    
                    numeric_matches1 = [match[0] for match in matches if match[0]]
                    word_matches1 = [match[1] for match in matches if match[1]]
                    
                    sum_totals = numeric_matches1 + word_matches1

                    print(sum_totals)


                    state = ""
                    for sum_total in sum_totals:
                        for total in totals:
                            if sum_total in total:
                                state = str("Experience {} years match " .format(total))
                                print(state)
                        

                    tokens = word_tokenize(skills.lower())
                    filtered_words = []
                    for token in tokens:
                        if token not in stop_words:
                            filtered_words.append(token)
                    print(filtered_words)
                    num_keywords = len(filtered_words)
                    matching_keywords = []
                    for word in filtered_words:
                        if word in text:
                            matching_keywords.append(word)



                    matching_keywords1 = len(matching_keywords)
                    print(matching_keywords)
                    matching_percentage = matching_keywords1 / num_keywords * 100
                    print(matching_percentage)
                    matching = []

                    if matching_percentage>=1:

                        if matching_percentage >= 75:
                            matching_files.append((file, matching_percentage,matching_keywords,state))
                        elif matching_percentage>=50 and matching_percentage<75:
                            matching_files1.append((file, matching_percentage,matching_keywords,state))
                        elif matching_percentage>=1 and matching_percentage<50:
                        # else:
                            matching_files2.append((file, matching_percentage,matching_keywords,state)) 
                        else:
                            print("No Files Found")

                    else:
                        print("No Files Found")
            
            sorted_files1 = sorted(matching_files, key=lambda x: x[1], reverse=True)

            sorted_files2 = sorted(matching_files1, key=lambda x: x[1], reverse=True)

            sorted_files3 = sorted(matching_files2, key=lambda x: x[1], reverse=True)
            
            filtered_files=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

            for file_info in sorted_files1:
                file_name = file_info[0]
                # file_path = folder_path + file_name
                filtered_files.append(file_name)

            print("filtered_files with sorted_files1 start")
            print(filtered_files)
            print("filtered_files with sorted_files1 end")




            filtered_files1=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

            for file_info in sorted_files2:
                file_name = file_info[0]
                # file_path = folder_path + file_name
                filtered_files1.append(file_name)

            print("filtered_files1 with sorted_files2 start")
            print(filtered_files1)
            print("filtered_files1 with sorted_files2 end")

            filtered_files2=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"


            resumes = []
            for file_info in sorted_files3:
                file_name = file_info[0]
                # file_path = folder_path + file_name

                filtered_files2.append(file_name)

            print("filtered_files2 with sorted_files3 start")
            print(filtered_files2)
            print("filtered_files2 with sorted_files3 start")
            

            matching_files_str = "<br><br>".join([f"{file}         {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state} " for file, percentage,matching_keywords,state in sorted_files1])
            matching_files_str1 = "<br><br>".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files2])
            matching_files_str2 = "<br><br>".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files3])
            

            count = len(matching_files + matching_files1 + matching_files2)

            print(f" {count} files were filtered ".format(count))
        elif experience:
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs"
            matching_files = []
            matching_files1 = []
            matching_files2 = []
            stop_words =stopwords.words("english")
            stop_words.extend(['.',",",":",")",",","software","(","want","person","knowledge","skills","percent","years","year","different","some", "experience","expertise","technology","technologies", "get", "must","resume","resumes","developer", "latest","engineer","exp","yrs",";","tech","experience"])
            files = os.listdir(folder_path)
            for file in files:
                if file.endswith('.docx'):
                    document = Document(os.path.join(folder_path, file))

                    text = "\n".join([paragraph.text for paragraph in document.paragraphs])
                    text = text.lower()


                    pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"
                    
                    matches = re.findall(pattern, text)
                    
                    numeric_matches = [match[0] for match in matches if match[0]]
                    word_matches = [match[1] for match in matches if match[1]]
                    
                    totals = numeric_matches + word_matches

                    print(totals)


                    summary_pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"

                    matches = re.findall(summary_pattern, experience)
                    
                    
                    numeric_matches1 = [match[0] for match in matches if match[0]]
                    word_matches1 = [match[1] for match in matches if match[1]]
                    
                    sum_totals = numeric_matches1 + word_matches1

                    print(sum_totals)


                    state = ""
                    for sum_total in sum_totals:
                        for total in totals:
                            if sum_total in total:
                                state = str("Experience {} years match " .format(total))
                                print(state)
                        

                    tokens = word_tokenize(skills.lower())
                    filtered_words = []
                    for token in tokens:
                        if token not in stop_words:
                            filtered_words.append(token)
                    print(filtered_words)
                    num_keywords = len(filtered_words)
                    matching_keywords = []
                    for word in filtered_words:
                        if word in text:
                            matching_keywords.append(word)



                    matching_keywords1 = len(matching_keywords)
                    print(matching_keywords)
                    matching_percentage = matching_keywords1 / num_keywords * 100
                    print(matching_percentage)
                    matching = []

                    if matching_percentage>=1:

                        if matching_percentage >= 75:
                            matching_files.append((file, matching_percentage,matching_keywords,state))
                        elif matching_percentage>=50 and matching_percentage<75:
                            matching_files1.append((file, matching_percentage,matching_keywords,state))
                        elif matching_percentage>=1 and matching_percentage<50:
                        # else:
                            matching_files2.append((file, matching_percentage,matching_keywords,state)) 
                        else:
                            print("No Files Found")

                    else:
                        print("No Files Found")
            
            sorted_files1 = sorted(matching_files, key=lambda x: x[1], reverse=True)

            sorted_files2 = sorted(matching_files1, key=lambda x: x[1], reverse=True)

            sorted_files3 = sorted(matching_files2, key=lambda x: x[1], reverse=True)
            
            filtered_files=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

            for file_info in sorted_files1:
                file_name = file_info[0]
                # file_path = folder_path + file_name
                filtered_files.append(file_name)

            print("filtered_files with sorted_files1 start")
            print(filtered_files)
            print("filtered_files with sorted_files1 end")




            filtered_files1=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

            for file_info in sorted_files2:
                file_name = file_info[0]
                # file_path = folder_path + file_name
                filtered_files1.append(file_name)

            print("filtered_files1 with sorted_files2 start")
            print(filtered_files1)
            print("filtered_files1 with sorted_files2 end")

            filtered_files2=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"


            resumes = []
            for file_info in sorted_files3:
                file_name = file_info[0]
                # file_path = folder_path + file_name

                filtered_files2.append(file_name)

            print("filtered_files2 with sorted_files3 start")
            print(filtered_files2)
            print("filtered_files2 with sorted_files3 start")
            

            matching_files_str = "<br><br>".join([f"{file}         {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state} " for file, percentage,matching_keywords,state in sorted_files1])
            matching_files_str1 = "<br><br>".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files2])
            matching_files_str2 = "<br><br>".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files3])
            

            count = len(matching_files + matching_files1 + matching_files2)

            print(f" {count} files were filtered ".format(count))
        elif skills:
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs"
            matching_files = []
            matching_files1 = []
            matching_files2 = []
            stop_words =stopwords.words("english")
            stop_words.extend(['.',",",":",")",",","software","(","want","person","knowledge","skills","percent","years","year","different","some", "experience","expertise","technology","technologies", "get", "must","resume","resumes","developer", "latest","engineer","exp","yrs",";","tech","experience"])
            files = os.listdir(folder_path)
            for file in files:
                if file.endswith('.docx'):
                    document = Document(os.path.join(folder_path, file))

                    text = "\n".join([paragraph.text for paragraph in document.paragraphs])
                    text = text.lower()


                    pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"
                    
                    matches = re.findall(pattern, text)
                    
                    numeric_matches = [match[0] for match in matches if match[0]]
                    word_matches = [match[1] for match in matches if match[1]]
                    
                    totals = numeric_matches + word_matches

                    print(totals)


                    summary_pattern = r"(?i)(?:(\d+(?:\.\d+)?|\d+\+?)\s*(?:years|yrs|yr.|years of experience))|(?:\b(\w+)\s+years\s+of\s+experience\b)"

                    matches = re.findall(summary_pattern, experience)
                    
                    
                    numeric_matches1 = [match[0] for match in matches if match[0]]
                    word_matches1 = [match[1] for match in matches if match[1]]
                    
                    sum_totals = numeric_matches1 + word_matches1

                    print(sum_totals)


                    state = ""
                    for sum_total in sum_totals:
                        for total in totals:
                            if sum_total in total:
                                state = str("Experience {} years match " .format(total))
                                print(state)
                        

                    tokens = word_tokenize(skills.lower())
                    filtered_words = []
                    for token in tokens:
                        if token not in stop_words:
                            filtered_words.append(token)
                    print(filtered_words)
                    num_keywords = len(filtered_words)
                    matching_keywords = []
                    for word in filtered_words:
                        if word in text:
                            matching_keywords.append(word)



                    matching_keywords1 = len(matching_keywords)
                    print(matching_keywords)
                    matching_percentage = matching_keywords1 / num_keywords * 100
                    print(matching_percentage)
                    matching = []

                    if matching_percentage>=1:

                        if matching_percentage >= 75:
                            matching_files.append((file, matching_percentage,matching_keywords,state))
                        elif matching_percentage>=50 and matching_percentage<75:
                            matching_files1.append((file, matching_percentage,matching_keywords,state))
                        elif matching_percentage>=1 and matching_percentage<50:
                        # else:
                            matching_files2.append((file, matching_percentage,matching_keywords,state)) 
                        else:
                            print("No Files Found")

                    else:
                        print("No Files Found")
            
            sorted_files1 = sorted(matching_files, key=lambda x: x[1], reverse=True)

            sorted_files2 = sorted(matching_files1, key=lambda x: x[1], reverse=True)

            sorted_files3 = sorted(matching_files2, key=lambda x: x[1], reverse=True)
            
            filtered_files=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

            for file_info in sorted_files1:
                file_name = file_info[0]
                # file_path = folder_path + file_name
                filtered_files.append(file_name)

            print("filtered_files with sorted_files1 start")
            print(filtered_files)
            print("filtered_files with sorted_files1 end")




            filtered_files1=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"

            for file_info in sorted_files2:
                file_name = file_info[0]
                # file_path = folder_path + file_name
                filtered_files1.append(file_name)

            print("filtered_files1 with sorted_files2 start")
            print(filtered_files1)
            print("filtered_files1 with sorted_files2 end")

            filtered_files2=[]
            folder_path = "C:/Users/sp13/OneDrive/Desktop/docs/"


            resumes = []
            for file_info in sorted_files3:
                file_name = file_info[0]
                # file_path = folder_path + file_name

                filtered_files2.append(file_name)

            print("filtered_files2 with sorted_files3 start")
            print(filtered_files2)
            print("filtered_files2 with sorted_files3 start")
            

            matching_files_str = "<br><br>".join([f"{file}         {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state} " for file, percentage,matching_keywords,state in sorted_files1])
            matching_files_str1 = "<br><br>".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files2])
            matching_files_str2 = "<br><br>".join([f"{file}        {percentage:.2f}%   with Matching Keywords  {matching_keywords} {state}" for file, percentage,matching_keywords,state in sorted_files3])
            

            count = len(matching_files + matching_files1 + matching_files2)

            print(f" {count} files were filtered ".format(count))
        context = {
            'count':count,
            'files1':filtered_files,
            'matching_files_str':matching_files_str,
            'files2':filtered_files1,
            'matching_files_str1':matching_files_str1,
            'files3':filtered_files2,
            'matching_files_str2':matching_files_str2,
            'experience':experience,
            'skills':skills,
        } 

        return render(request,'resume-finder-classic.html',context)
    else:
        return render(request,'resume-finder-classic.html')



